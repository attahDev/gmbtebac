import io
import logging
from datetime import datetime
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas as rl_canvas
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from services.template_registry import get_template, get_variation
logger = logging.getLogger(__name__)
PAGE_W, PAGE_H = A4
DEFAULT_PRIMARY = '#1A3557'
DEFAULT_SECONDARY = '#E87722'

def _hex(h: str) -> HexColor:
    return HexColor(h if h else DEFAULT_PRIMARY)

def _primary(inputs: dict) -> HexColor:
    return _hex(inputs.get('primary_color') or DEFAULT_PRIMARY)

def _secondary(inputs: dict) -> HexColor:
    return _hex(inputs.get('secondary_color') or DEFAULT_SECONDARY)

def _rgb(h: str):
    h = (h or DEFAULT_PRIMARY).lstrip('#')
    return tuple((int(h[i:i + 2], 16) / 255 for i in (0, 2, 4)))

def _lum(h: str) -> float:
    r, g, b = _rgb(h)
    return 0.299 * r + 0.587 * g + 0.114 * b

def _on_dark(h: str) -> HexColor:
    return HexColor('#FFFFFF') if _lum(h) < 0.55 else HexColor('#111111')

def _styles(primary: str, secondary: str) -> dict:
    base = getSampleStyleSheet()
    p = primary or DEFAULT_PRIMARY
    s = secondary or DEFAULT_SECONDARY
    return {'body': ParagraphStyle('body', parent=base['Normal'], fontSize=9, leading=13, textColor=HexColor('#333333')), 'h1': ParagraphStyle('h1', parent=base['Normal'], fontSize=16, leading=20, textColor=_hex(p), fontName='Helvetica-Bold'), 'h2': ParagraphStyle('h2', parent=base['Normal'], fontSize=11, leading=15, textColor=_hex(p), fontName='Helvetica-Bold'), 'h3': ParagraphStyle('h3', parent=base['Normal'], fontSize=9, leading=13, textColor=_hex(p), fontName='Helvetica-Bold'), 'label': ParagraphStyle('label', parent=base['Normal'], fontSize=8, leading=11, textColor=HexColor('#777777'), fontName='Helvetica'), 'accent': ParagraphStyle('accent', parent=base['Normal'], fontSize=9, leading=13, textColor=_hex(s), fontName='Helvetica-Bold'), 'center': ParagraphStyle('center', parent=base['Normal'], fontSize=9, leading=13, alignment=TA_CENTER), 'right': ParagraphStyle('right', parent=base['Normal'], fontSize=9, leading=13, alignment=TA_RIGHT), 'small': ParagraphStyle('small', parent=base['Normal'], fontSize=7.5, leading=11, textColor=HexColor('#555555')), 'white': ParagraphStyle('white', parent=base['Normal'], fontSize=9, leading=13, textColor=HexColor('#FFFFFF'))}

def _draw_rect(c, x, y, w, h, fill_hex):
    c.setFillColor(_hex(fill_hex))
    c.rect(x, y, w, h, fill=1, stroke=0)

def _draw_polygon(c, pts, fill_hex):
    c.setFillColor(_hex(fill_hex))
    path = c.beginPath()
    path.moveTo(*pts[0])
    for pt in pts[1:]:
        path.lineTo(*pt)
    path.close()
    c.drawPath(path, fill=1, stroke=0)

def _draw_line(c, x1, y1, x2, y2, color_hex, width=1):
    c.setStrokeColor(_hex(color_hex))
    c.setLineWidth(width)
    c.line(x1, y1, x2, y2)

def _draw_text_on_canvas(c, text, x, y, font='Helvetica', size=8, color='#333333'):
    c.setFont(font, size)
    c.setFillColor(_hex(color))
    c.drawString(x, y, text)

def _draw_watermark(c, text, page_w, page_h, color_hex, alpha=0.05):
    c.saveState()
    c.setFillColor(_hex(color_hex))
    c.setFillAlpha(alpha)
    c.setFont('Helvetica-Bold', 80)
    c.translate(page_w / 2, page_h / 2)
    c.rotate(30)
    c.drawCentredString(0, 0, text)
    c.restoreState()

def _lh_corporate(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = (36 + var['header_extra_mm']) * mm
    tilt = var['header_tilt_pt']
    block_w = pw * 0.58
    _draw_rect(c, 0, ph - hh, block_w, hh, pri)
    _draw_polygon(c, [(block_w, ph - hh), (block_w + tilt, ph), (block_w - tilt / 2, ph), (block_w - tilt, ph - hh)], pri)
    _draw_polygon(c, [(block_w + tilt, ph), (block_w + tilt * 2.5, ph), (block_w + tilt * 1.8, ph - hh), (block_w + tilt * 0.2, ph - hh)], sec)
    company = inputs.get('company_name', '')
    tagline = inputs.get('tagline', '')
    c.saveState()
    c.setFillColor(HexColor('#FFFFFF'))
    c.setFont('Helvetica-Bold', 14)
    c.drawString(14 * mm, ph - hh / 2 - 4, company)
    if tagline:
        c.setFont('Helvetica', 8)
        c.setFillColor(HexColor('#DDDDDD'))
        c.drawString(14 * mm, ph - hh / 2 - 14, tagline)
    c.restoreState()
    _lh_draw_logo(c, inputs, x=14 * mm, y=ph - hh + 4 * mm, max_w=40 * mm, max_h=hh - 8 * mm)
    _lh_draw_contact_right(c, inputs, pw, ph, hh, sec, '#FFFFFF')
    fh = (14 + var['footer_extra_mm']) * mm
    _draw_rect(c, 0, 0, pw * 0.35, fh, pri)
    _draw_rect(c, pw * 0.4, 0, pw * 0.6, fh * 0.5, sec)

def _lh_minimal(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = (26 + var['header_extra_mm']) * mm
    _draw_rect(c, 0, ph - 3, pw, 3, sec)
    c.saveState()
    c.setFillColor(_hex(pri))
    if inputs.get('logo_url'):
        _lh_draw_logo(c, inputs, x=18 * mm, y=ph - hh + 4 * mm, max_w=45 * mm, max_h=hh - 8 * mm)
    else:
        c.setFont('Helvetica-Bold', 13)
        c.drawString(18 * mm, ph - hh / 2, inputs.get('company_name', ''))
    if inputs.get('tagline'):
        c.setFont('Helvetica', 8)
        c.setFillColor(HexColor('#888888'))
        c.drawString(18 * mm, ph - hh / 2 - 12, inputs.get('tagline', ''))
    c.restoreState()
    _lh_draw_contact_right(c, inputs, pw, ph, hh, sec, pri)
    rule_y = ph - hh - 2 * mm
    _draw_line(c, 18 * mm, rule_y, pw - 18 * mm, rule_y, pri, var['rule_thickness'])
    _draw_line(c, 18 * mm, 14 * mm, pw - 18 * mm, 14 * mm, pri, var['rule_thickness'])
    c.saveState()
    c.setFont('Helvetica', 7)
    c.setFillColor(HexColor('#999999'))
    website = inputs.get('website', '')
    if website:
        c.drawCentredString(pw / 2, 9 * mm, website)
    c.restoreState()

def _lh_dual_rule(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = (28 + var['header_extra_mm']) * mm
    company = inputs.get('company_name', '')
    rc = inputs.get('registration_number', '')
    c.saveState()
    c.setFillColor(_hex(pri))
    if inputs.get('logo_url'):
        _lh_draw_logo(c, inputs, x=18 * mm, y=ph - hh + 4 * mm, max_w=50 * mm, max_h=hh - 8 * mm)
    else:
        c.setFont('Helvetica-Bold', 14)
        c.drawString(18 * mm, ph - hh / 2, company)
    if rc:
        c.setFont('Helvetica', 7.5)
        c.setFillColor(HexColor('#888888'))
        c.drawString(18 * mm, ph - hh / 2 - 13, rc)
    c.restoreState()
    _lh_draw_contact_right(c, inputs, pw, ph, hh, sec, pri)
    rule_thick = var['rule_thickness'] + 1
    rule_thin = var['rule_thickness']
    rule_y1 = ph - hh - 2 * mm
    rule_y2 = rule_y1 - 3
    _draw_rect(c, 18 * mm, rule_y1, pw - 36 * mm, rule_thick + 1, pri)
    _draw_rect(c, 18 * mm, rule_y2 - rule_thin, pw - 36 * mm, rule_thin, sec)
    if inputs.get('watermark_logo'):
        _draw_watermark(c, company[:2].upper(), pw, ph, pri, alpha=0.04)
    fh = (12 + var['footer_extra_mm']) * mm
    _draw_polygon(c, [(0, 0), (0, fh), (pw * 0.32, fh), (pw * 0.28, 0)], pri)
    _draw_polygon(c, [(pw * 0.3, 0), (pw * 0.35, fh * 0.6), (pw, fh * 0.6), (pw, 0)], sec)

def _lh_wave(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = (42 + var['header_extra_mm']) * mm
    dip = 14 * mm
    _draw_polygon(c, [(0, ph), (pw, ph), (pw, ph - hh + dip), (pw * 0.75, ph - hh), (pw * 0.5, ph - hh + dip / 2), (pw * 0.25, ph - hh), (0, ph - hh + dip)], pri)
    _draw_polygon(c, [(0, ph - hh + dip - 5), (pw * 0.25, ph - hh - 5), (pw * 0.5, ph - hh + dip / 2 - 5), (pw * 0.75, ph - hh - 5), (pw, ph - hh + dip - 5), (pw, ph - hh + dip - 5 + 6), (pw * 0.75, ph - hh - 5 + 6), (pw * 0.5, ph - hh + dip / 2 - 5 + 6), (pw * 0.25, ph - hh - 5 + 6), (0, ph - hh + dip - 5 + 6)], sec)
    c.saveState()
    c.setFillColor(HexColor('#FFFFFF'))
    if inputs.get('logo_url'):
        _lh_draw_logo(c, inputs, x=16 * mm, y=ph - hh + 6 * mm, max_w=50 * mm, max_h=hh - 14 * mm)
    else:
        c.setFont('Helvetica-Bold', 14)
        c.drawString(16 * mm, ph - hh / 2 + 6, inputs.get('company_name', ''))
    if inputs.get('tagline'):
        c.setFont('Helvetica', 8)
        c.setFillColor(HexColor('#CCDDEE'))
        c.drawString(16 * mm, ph - hh / 2 - 6, inputs.get('tagline', ''))
    c.restoreState()
    _lh_draw_contact_right(c, inputs, pw, ph, hh, sec, '#FFFFFF')
    fh = (18 + var['footer_extra_mm']) * mm
    _draw_polygon(c, [(0, 0), (pw, 0), (pw, fh - dip), (pw * 0.75, fh), (pw * 0.5, fh - dip / 2), (pw * 0.25, fh), (0, fh - dip)], pri)

def _lh_draw_contact_right(c, inputs, pw, ph, hh, sec_hex, text_hex):
    items = []
    if inputs.get('phone'):
        items.append(('☎', inputs['phone']))
    if inputs.get('email'):
        items.append(('✉', inputs['email']))
    if inputs.get('website'):
        items.append(('⊕', inputs['website']))
    c.saveState()
    y = ph - hh / 2 + (len(items) - 1) * 8
    for icon, val in items:
        c.setFont('Helvetica-Bold', 8)
        c.setFillColor(_hex(sec_hex))
        c.drawRightString(pw - 38 * mm, y, icon)
        c.setFont('Helvetica', 8)
        c.setFillColor(_hex(text_hex if text_hex != sec_hex else '#333333'))
        c.drawRightString(pw - 18 * mm, y, val)
        y -= 12
    c.restoreState()

def _lh_draw_logo(c, inputs, x, y, max_w=45 * mm, max_h=18 * mm):
    logo_url = inputs.get('logo_url')
    if not logo_url:
        return
    try:
        import httpx
        from reportlab.lib.utils import ImageReader
        resp = httpx.get(logo_url, timeout=8, follow_redirects=True)
        resp.raise_for_status()
        from PIL import Image as PILImage
        pil = PILImage.open(io.BytesIO(resp.content)).convert('RGBA')
        from reportlab.lib.units import mm as rl_mm
        img_w_pt = pil.width * (72 / 96)
        img_h_pt = pil.height * (72 / 96)
        scale = min(max_w / img_w_pt, max_h / img_h_pt, 1.0)
        draw_w = img_w_pt * scale
        draw_h = img_h_pt * scale
        buf = io.BytesIO()
        pil.save(buf, format='PNG')
        buf.seek(0)
        c.drawImage(ImageReader(buf), x, y, width=draw_w, height=draw_h, mask='auto')
    except Exception as e:
        logger.warning(f'Letterhead logo draw failed: {e}')
_LH_DRAWERS = {'corporate': _lh_corporate, 'minimal': _lh_minimal, 'dual_rule': _lh_dual_rule, 'wave': _lh_wave}

def _inv_diagonal_split(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = 42 * mm
    tilt = var['header_tilt_pt'] * 1.4
    _draw_rect(c, 0, ph - hh, pw * 0.62, hh, pri)
    _draw_polygon(c, [(pw * 0.62, ph - hh), (pw * 0.62 + tilt, ph), (pw * 0.78 + tilt, ph), (pw * 0.78, ph - hh)], sec)
    _draw_rect(c, pw * 0.78 + tilt, ph - hh, pw, hh, sec)
    c.saveState()
    c.setFillColor(HexColor('#FFFFFF'))
    c.setFont('Helvetica-Bold', 16)
    c.drawString(14 * mm, ph - hh / 2 + 2, inputs.get('company_name', ''))
    c.setFont('Helvetica', 8)
    c.setFillColor(HexColor('#CCDDEE'))
    c.drawString(14 * mm, ph - hh / 2 - 11, inputs.get('company_address', ''))
    c.restoreState()
    c.saveState()
    c.setFillColor(_on_dark(sec))
    c.setFont('Helvetica-Bold', 22)
    c.drawRightString(pw - 12 * mm, ph - hh / 2 - 6, 'INVOICE')
    c.restoreState()
    _draw_rect(c, 0, 0, pw, 14 * mm, pri)
    c.saveState()
    c.setFont('Helvetica', 8)
    c.setFillColor(HexColor('#FFFFFF'))
    c.drawCentredString(pw / 2, 5 * mm, 'Thank You For Your Business')
    c.restoreState()

def _inv_logo_title_rule(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = 32 * mm
    c.saveState()
    c.setFillColor(_hex(pri))
    c.setFont('Helvetica-Bold', 14)
    c.drawString(18 * mm, ph - hh / 2 + 4, inputs.get('company_name', ''))
    c.setFont('Helvetica', 8)
    c.setFillColor(HexColor('#666666'))
    c.drawString(18 * mm, ph - hh / 2 - 8, inputs.get('company_address', ''))
    c.setFont('Helvetica', 8)
    c.drawString(18 * mm, ph - hh / 2 - 19, inputs.get('email', ''))
    c.restoreState()
    c.saveState()
    c.setFillColor(_hex(pri))
    c.setFont('Helvetica-Bold', 28)
    c.drawRightString(pw - 18 * mm, ph - hh / 2, 'INVOICE')
    c.restoreState()
    rule_y = ph - hh - 2 * mm
    _draw_rect(c, 18 * mm, rule_y - 2, pw - 36 * mm, var['rule_thickness'] + 1, pri)
    _draw_rect(c, 18 * mm, rule_y - 4, pw - 36 * mm, 2, sec)
    fh = 16 * mm
    _draw_rect(c, 0, 0, pw, fh, pri)
    c.saveState()
    c.setFont('Helvetica', 7.5)
    c.setFillColor(HexColor('#AABBCC'))
    parts = []
    if inputs.get('website'):
        parts.append(inputs['website'])
    if inputs.get('phone'):
        parts.append(inputs['phone'])
    if inputs.get('email'):
        parts.append(inputs['email'])
    c.drawCentredString(pw / 2, fh / 2 - 3, '  |  '.join(parts))
    c.restoreState()

def _inv_text_only(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    c.saveState()
    c.setFillColor(_hex(pri))
    c.setFont('Helvetica-Bold', 26)
    c.drawString(18 * mm, ph - 28 * mm, 'INVOICE')
    c.setFont('Helvetica', 8.5)
    c.setFillColor(HexColor('#444444'))
    c.drawString(18 * mm, ph - 38 * mm, inputs.get('company_name', ''))
    c.drawString(18 * mm, ph - 48 * mm, inputs.get('company_address', ''))
    c.restoreState()
    _draw_line(c, 18 * mm, ph - 56 * mm, pw - 18 * mm, ph - 56 * mm, '#CCCCCC', 0.5)
    _draw_line(c, 18 * mm, 18 * mm, pw - 18 * mm, 18 * mm, '#CCCCCC', 0.5)
_INV_DRAWERS = {'bold': _inv_diagonal_split, 'professional': _inv_logo_title_rule, 'minimal': _inv_text_only}

def _qt_centered_title(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    _draw_rect(c, 18 * mm, ph - 12 * mm, pw - 36 * mm, 2.5, sec)
    c.saveState()
    c.setFillColor(_hex(pri))
    c.setFont('Helvetica-Bold', 12)
    c.drawString(18 * mm, ph - 22 * mm, inputs.get('company_name', ''))
    c.setFont('Helvetica', 8)
    c.setFillColor(HexColor('#555555'))
    c.drawString(18 * mm, ph - 32 * mm, inputs.get('company_address', ''))
    rc = inputs.get('registration_number', '')
    if rc:
        c.setFont('Helvetica', 7.5)
        c.setFillColor(HexColor('#888888'))
        c.drawString(18 * mm, ph - 40 * mm, rc)
    c.restoreState()
    fh = 14 * mm
    _draw_rect(c, 0, 0, pw, fh, pri)
    c.saveState()
    c.setFont('Helvetica', 7.5)
    c.setFillColor(HexColor('#AABBCC'))
    parts = []
    if inputs.get('phone'):
        parts.append(inputs['phone'])
    if inputs.get('email'):
        parts.append(inputs['email'])
    if inputs.get('website'):
        parts.append(inputs['website'])
    c.drawCentredString(pw / 2, fh / 2 - 3, '  |  '.join(parts))
    c.restoreState()

def _qt_split_header(c, inputs, var, pw, ph):
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    hh = 30 * mm
    _draw_rect(c, 0, ph - hh, 5, hh, sec)
    c.saveState()
    c.setFillColor(_hex(pri))
    c.setFont('Helvetica-Bold', 13)
    c.drawString(12 * mm, ph - hh / 2 + 4, inputs.get('company_name', ''))
    c.setFont('Helvetica', 7.5)
    c.setFillColor(HexColor('#888888'))
    c.drawString(12 * mm, ph - hh / 2 - 7, inputs.get('company_address', ''))
    rc = inputs.get('registration_number', '')
    if rc:
        c.setFont('Helvetica', 7)
        c.setFillColor(HexColor('#AAAAAA'))
        c.drawString(12 * mm, ph - hh / 2 - 17, rc)
    c.restoreState()
    c.saveState()
    c.setFillColor(_hex(sec))
    c.setFont('Helvetica-BoldOblique', 22)
    c.drawRightString(pw - 18 * mm, ph - hh / 2 - 4, 'Quote')
    c.restoreState()
    rule_y = ph - hh - 3 * mm
    _draw_rect(c, 12 * mm, rule_y, pw - 24 * mm, 1.5, pri)
    fh = 16 * mm
    _draw_rect(c, 0, 0, pw, fh, pri)
    c.saveState()
    c.setFont('Helvetica-Bold', 8)
    c.setFillColor(HexColor('#FFFFFF'))
    c.drawCentredString(pw / 2, fh / 2 + 2, 'THANK YOU FOR YOUR BUSINESS!')
    c.setFont('Helvetica', 7)
    c.setFillColor(HexColor('#AABBCC'))
    parts = []
    if inputs.get('phone'):
        parts.append(inputs['phone'])
    if inputs.get('email'):
        parts.append(inputs['email'])
    if inputs.get('website'):
        parts.append(inputs['website'])
    c.drawCentredString(pw / 2, fh / 2 - 7, '  |  '.join(parts))
    c.restoreState()
_QT_DRAWERS = {'formal': _qt_centered_title, 'modern': _qt_split_header}

def _make_doc(buf, top_mm, bottom_mm, left_mm, right_mm, bg_fn, inputs, var):
    pw, ph = (PAGE_W, PAGE_H)
    lm = left_mm * mm
    rm = right_mm * mm
    tm = top_mm * mm
    bm = bottom_mm * mm
    frame = Frame(lm, bm, pw - lm - rm, ph - tm - bm, id='main', showBoundary=0)

    def on_page(c, doc):
        c.saveState()
        bg_fn(c, inputs, var, pw, ph)
        c.restoreState()
    pt = PageTemplate(id='main', frames=[frame], onPage=on_page)
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=lm, rightMargin=rm, topMargin=tm, bottomMargin=bm)
    doc.addPageTemplates([pt])
    return doc

def _items_table(items, currency_sym, primary_hex, style_key, var):
    sym = currency_sym.split()[-1] if currency_sym else '₦'
    header = [Paragraph('<b>#</b>', ParagraphStyle('th', fontSize=8, textColor=HexColor('#FFFFFF'), fontName='Helvetica-Bold')), Paragraph('<b>Description</b>', ParagraphStyle('th', fontSize=8, textColor=HexColor('#FFFFFF'), fontName='Helvetica-Bold')), Paragraph('<b>Qty</b>', ParagraphStyle('th', fontSize=8, textColor=HexColor('#FFFFFF'), fontName='Helvetica-Bold', alignment=TA_CENTER)), Paragraph('<b>Unit Price</b>', ParagraphStyle('th', fontSize=8, textColor=HexColor('#FFFFFF'), fontName='Helvetica-Bold', alignment=TA_RIGHT)), Paragraph('<b>Total</b>', ParagraphStyle('th', fontSize=8, textColor=HexColor('#FFFFFF'), fontName='Helvetica-Bold', alignment=TA_RIGHT))]
    rows = [header]
    p_body = ParagraphStyle('td', fontSize=8, leading=11)
    p_right = ParagraphStyle('tdr', fontSize=8, leading=11, alignment=TA_RIGHT)
    p_center = ParagraphStyle('tdc', fontSize=8, leading=11, alignment=TA_CENTER)
    for i, item in enumerate(items or []):
        desc = item.get('description', '')
        qty = item.get('qty', item.get('quantity', 1))
        price = float(item.get('unit_price', item.get('price', 0)))
        total = float(item.get('total', qty * price))
        rows.append([Paragraph(str(i + 1), p_center), Paragraph(desc, p_body), Paragraph(str(qty), p_center), Paragraph(f'{sym}{price:,.2f}', p_right), Paragraph(f'{sym}{total:,.2f}', p_right)])
    col_widths = [10 * mm, 78 * mm, 18 * mm, 32 * mm, 32 * mm]
    if style_key in ('dark_bg', 'primary_bg'):
        hdr_bg = HexColor(primary_hex)
        ts = TableStyle([('BACKGROUND', (0, 0), (-1, 0), hdr_bg), ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#FFFFFF'), HexColor('#F5F6F8')]), ('GRID', (0, 0), (-1, -1), 0.3, HexColor('#DDDDDD')), ('TOPPADDING', (0, 0), (-1, -1), var['table_row_pad_pt']), ('BOTTOMPADDING', (0, 0), (-1, -1), var['table_row_pad_pt']), ('LEFTPADDING', (0, 0), (-1, -1), 4), ('RIGHTPADDING', (0, 0), (-1, -1), 4), ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')])
    else:
        ts = TableStyle([('LINEBELOW', (0, 0), (-1, 0), 1.5, HexColor(primary_hex)), ('LINEBELOW', (0, 1), (-1, -1), 0.3, HexColor('#DDDDDD')), ('TEXTCOLOR', (0, 0), (-1, 0), HexColor(primary_hex)), ('TOPPADDING', (0, 0), (-1, -1), var['table_row_pad_pt']), ('BOTTOMPADDING', (0, 0), (-1, -1), var['table_row_pad_pt']), ('LEFTPADDING', (0, 0), (-1, -1), 4), ('RIGHTPADDING', (0, 0), (-1, -1), 4), ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')])
    return Table(rows, colWidths=col_widths, style=ts, repeatRows=1)

def _totals_table(subtotal, tax_rate, discount, currency_sym, primary_hex):
    sym = currency_sym.split()[-1] if currency_sym else '₦'
    pri = HexColor(primary_hex)
    rows = []
    p_lbl = ParagraphStyle('tl', fontSize=8.5, alignment=TA_RIGHT)
    p_val = ParagraphStyle('tv', fontSize=8.5, alignment=TA_RIGHT, fontName='Helvetica-Bold')
    rows.append([Paragraph('Subtotal:', p_lbl), Paragraph(f'{sym}{subtotal:,.2f}', p_val)])
    if discount and discount > 0:
        rows.append([Paragraph('Discount:', p_lbl), Paragraph(f'-{sym}{discount:,.2f}', p_val)])
        subtotal -= discount
    if tax_rate and tax_rate > 0:
        tax_amt = subtotal * tax_rate / 100
        rows.append([Paragraph(f'Tax ({tax_rate:.1f}%):', p_lbl), Paragraph(f'{sym}{tax_amt:,.2f}', p_val)])
        grand = subtotal + tax_amt
    else:
        grand = subtotal
    p_grand_lbl = ParagraphStyle('gl', fontSize=10, alignment=TA_RIGHT, fontName='Helvetica-Bold', textColor=pri)
    p_grand_val = ParagraphStyle('gv', fontSize=10, alignment=TA_RIGHT, fontName='Helvetica-Bold', textColor=pri)
    rows.append([Paragraph('Total:', p_grand_lbl), Paragraph(f'{sym}{grand:,.2f}', p_grand_val)])
    t = Table(rows, colWidths=[50 * mm, 38 * mm])
    t.setStyle(TableStyle([('LINEABOVE', (0, -1), (-1, -1), 1.5, pri), ('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3), ('ALIGN', (0, 0), (-1, -1), 'RIGHT')]))
    return t

def _badge(text, bg_hex, fg_hex='#FFFFFF'):
    style = ParagraphStyle('badge', fontSize=8, textColor=HexColor(fg_hex), fontName='Helvetica-Bold', backColor=HexColor(bg_hex), borderPadding=(2, 5, 2, 5))
    return Paragraph(text, style)

def build_letterhead_pdf(inputs: dict, asset_id: str='default') -> bytes:
    variant = inputs.get('template_variant', 'corporate')
    tmpl = get_template('letterhead', variant)
    var = get_variation(asset_id)
    bg_fn = _LH_DRAWERS.get(variant, _lh_corporate)
    buf = io.BytesIO()
    doc = _make_doc(buf, top_mm=tmpl['body_top_mm'] + var['header_extra_mm'], bottom_mm=tmpl['body_bottom_mm'] + var['footer_extra_mm'], left_mm=tmpl['body_left_mm'], right_mm=tmpl['body_right_mm'], bg_fn=bg_fn, inputs=inputs, var=var)
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)
    story = [Spacer(1, 4 * mm), Paragraph('Dear Sir / Madam,', sty['body']), Spacer(1, 4 * mm), Paragraph('This letterhead is ready for use. Replace this placeholder content with your correspondence text. The header and footer on this document carry your brand identity automatically.', sty['body']), Spacer(1, 6 * mm), Paragraph('Yours faithfully,', sty['body']), Spacer(1, 12 * mm), Paragraph(inputs.get('company_name', ''), sty['h3'])]
    doc.build(story)
    return buf.getvalue()

def build_invoice_pdf(inputs: dict, asset_id: str='default') -> bytes:
    variant = inputs.get('template_variant', 'professional')
    tmpl = get_template('invoice', variant)
    var = get_variation(asset_id)
    bg_fn = _INV_DRAWERS.get(variant, _inv_logo_title_rule)
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)
    sym = (inputs.get('currency') or 'NGN ₦').split()[-1]
    buf = io.BytesIO()
    doc = _make_doc(buf, top_mm=tmpl['body_top_mm'], bottom_mm=tmpl['body_bottom_mm'], left_mm=20, right_mm=20, bg_fn=bg_fn, inputs=inputs, var=var)
    today = datetime.now().strftime('%d %B %Y')
    inv_prefix = inputs.get('invoice_number_prefix', 'INV-')
    story = []
    if tmpl.get('bill_badge'):
        story.append(_badge('  Invoice To  ', pri))
    else:
        story.append(Paragraph('<b>INVOICE TO:</b>', sty['h3']))
    story.append(Spacer(1, 2 * mm))
    client = inputs.get('client_name', '[Client Name]')
    story.append(Paragraph(f'<b>{client}</b>', sty['body']))
    for line in [inputs.get('client_address', ''), inputs.get('client_email', ''), inputs.get('client_phone', '')]:
        if line:
            story.append(Paragraph(line, sty['body']))
    story.append(Spacer(1, 3 * mm))
    meta_data = [[Paragraph('<b>Invoice #:</b>', sty['label']), Paragraph(f"{inv_prefix}{inputs.get('invoice_number', '001')}", sty['body'])], [Paragraph('<b>Date:</b>', sty['label']), Paragraph(inputs.get('invoice_date', today), sty['body'])], [Paragraph('<b>Due Date:</b>', sty['label']), Paragraph(inputs.get('due_date', ''), sty['body'])]]
    if inputs.get('payment_terms'):
        meta_data.append([Paragraph('<b>Terms:</b>', sty['label']), Paragraph(inputs['payment_terms'], sty['body'])])
    meta_t = Table(meta_data, colWidths=[28 * mm, 60 * mm])
    meta_t.setStyle(TableStyle([('TOPPADDING', (0, 0), (-1, -1), 2), ('BOTTOMPADDING', (0, 0), (-1, -1), 2), ('LEFTPADDING', (0, 0), (-1, -1), 0)]))
    story.append(meta_t)
    story.append(Spacer(1, 4 * mm))
    items = inputs.get('items', [])
    story.append(_items_table(items, sym, pri, tmpl.get('table_header', 'primary_bg'), var))
    story.append(Spacer(1, 3 * mm))
    subtotal = sum((float(it.get('total', float(it.get('qty', it.get('quantity', 1))) * float(it.get('unit_price', it.get('price', 0))))) for it in items))
    tax_rate = float(inputs.get('tax_rate', 0))
    discount = inputs.get('discount')
    totals_t = _totals_table(subtotal, tax_rate, discount, sym, pri)
    totals_wrap = Table([[totals_t]], colWidths=[170 * mm])
    totals_wrap.setStyle(TableStyle([('ALIGN', (0, 0), (0, 0), 'RIGHT')]))
    story.append(totals_wrap)
    story.append(Spacer(1, 5 * mm))
    if tmpl.get('payment_section') and inputs.get('payment_methods'):
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#CCCCCC')))
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph('<b>PAYMENT METHODS</b>', sty['h3']))
        story.append(Spacer(1, 2 * mm))
        for pm in inputs['payment_methods']:
            story.append(Paragraph(f"<b>{pm.get('label', '')}:</b>  {pm.get('bank_name', '')}  {pm.get('account_name', '')}  {pm.get('account_number', '')}", sty['small']))
            if pm.get('details'):
                story.append(Paragraph(pm['details'], sty['small']))
        story.append(Spacer(1, 3 * mm))
    if inputs.get('footer_note'):
        story.append(Paragraph('<b>Notes:</b>', sty['h3']))
        story.append(Paragraph(inputs['footer_note'], sty['small']))
        story.append(Spacer(1, 3 * mm))
    if tmpl.get('terms_section') and inputs.get('terms_and_conditions'):
        story.append(Paragraph('<b>Terms & Conditions:</b>', sty['h3']))
        story.append(Paragraph(inputs['terms_and_conditions'], sty['small']))
    doc.build(story)
    return buf.getvalue()

def build_quotation_pdf(inputs: dict, asset_id: str='default') -> bytes:
    variant = inputs.get('template_variant', 'modern')
    tmpl = get_template('quotation', variant)
    var = get_variation(asset_id)
    bg_fn = _QT_DRAWERS.get(variant, _qt_split_header)
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)
    sym = (inputs.get('currency') or 'NGN ₦').split()[-1]
    buf = io.BytesIO()
    doc = _make_doc(buf, top_mm=tmpl['body_top_mm'], bottom_mm=tmpl['body_bottom_mm'], left_mm=20, right_mm=20, bg_fn=bg_fn, inputs=inputs, var=var)
    today = datetime.now().strftime('%d %B %Y')
    qt_num = f"{inputs.get('quote_number_prefix', 'QT-')}{inputs.get('quote_number', '001')}"
    exp_date = inputs.get('expiration_date', '')
    story = []
    if tmpl['header_style'] == 'centered_title':
        story.append(Paragraph('FORMAL QUOTATION', ParagraphStyle('fq', fontSize=14, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=_hex(pri))))
        story.append(Spacer(1, 4 * mm))
    meta = [[Paragraph('<b>Date:</b>', sty['label']), Paragraph(today, sty['body'])], [Paragraph('<b>Quote No:</b>', sty['label']), Paragraph(qt_num, sty['body'])], [Paragraph('<b>Valid For:</b>', sty['label']), Paragraph(inputs.get('quote_valid_for', '30 days'), sty['body'])]]
    if exp_date:
        meta.append([Paragraph('<b>Expires:</b>', sty['label']), Paragraph(exp_date, sty['body'])])
    if inputs.get('prepared_by'):
        meta.append([Paragraph('<b>Prepared By:</b>', sty['label']), Paragraph(inputs['prepared_by'], sty['body'])])
    if inputs.get('registration_number'):
        meta.append([Paragraph('<b>Reg No:</b>', sty['label']), Paragraph(inputs['registration_number'], sty['body'])])
    meta_t = Table(meta, colWidths=[32 * mm, 70 * mm])
    meta_t.setStyle(TableStyle([('TOPPADDING', (0, 0), (-1, -1), 2), ('BOTTOMPADDING', (0, 0), (-1, -1), 2), ('LEFTPADDING', (0, 0), (-1, -1), 0)]))
    story.append(meta_t)
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph('<b>Quotation For:</b>', sty['h3']))
    for line in [inputs.get('client_name', '[Client Name]'), inputs.get('client_address', ''), inputs.get('client_phone', ''), inputs.get('client_email', '')]:
        if line:
            story.append(Paragraph(line, sty['body']))
    story.append(Spacer(1, 4 * mm))
    items = inputs.get('items', [])
    story.append(_items_table(items, sym, pri, 'primary_bg', var))
    story.append(Spacer(1, 3 * mm))
    subtotal = sum((float(it.get('total', float(it.get('qty', it.get('quantity', 1))) * float(it.get('unit_price', it.get('price', 0))))) for it in items))
    tax_rate = float(inputs.get('tax_rate', inputs.get('vat_rate', 0)))
    totals_t = _totals_table(subtotal, tax_rate, None, sym, pri)
    totals_wrap = Table([[totals_t]], colWidths=[170 * mm])
    totals_wrap.setStyle(TableStyle([('ALIGN', (0, 0), (0, 0), 'RIGHT')]))
    story.append(totals_wrap)
    story.append(Spacer(1, 4 * mm))
    if tmpl.get('show_delivery') and (inputs.get('delivery_required') is not None or inputs.get('packaging_required') is not None):
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#CCCCCC')))
        story.append(Spacer(1, 2 * mm))
        dr = inputs.get('delivery_required')
        pr = inputs.get('packaging_required')
        if dr is not None:
            story.append(Paragraph(f"Items to be delivered?   [{('X' if dr else ' ')}] Yes   [{('X' if not dr else ' ')}] No", sty['body']))
        if pr is not None:
            story.append(Paragraph(f"Packaging required?      [{('X' if pr else ' ')}] Yes   [{('X' if not pr else ' ')}] No", sty['body']))
        story.append(Spacer(1, 3 * mm))
    if tmpl.get('terms_section') and inputs.get('terms_and_conditions'):
        story.append(Paragraph('<b>Terms and Conditions:</b>', sty['h3']))
        story.append(Spacer(1, 1 * mm))
        for i, line in enumerate(inputs['terms_and_conditions'].split('\n'), 1):
            if line.strip():
                story.append(Paragraph(f'{i}. {line.strip()}', sty['small']))
        story.append(Spacer(1, 3 * mm))
    if inputs.get('payment_terms'):
        story.append(Paragraph('<b>Payment Terms:</b>', sty['h3']))
        story.append(Paragraph(inputs['payment_terms'], sty['small']))
        story.append(Spacer(1, 3 * mm))
    if tmpl.get('show_signature') and inputs.get('signature_section', True):
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#CCCCCC')))
        story.append(Spacer(1, 3 * mm))
        sig_data = [[Paragraph('<b>Name</b>', sty['label']), Paragraph('<b>Signature</b>', sty['label']), Paragraph('<b>Date</b>', sty['label'])], [Paragraph('_' * 30, sty['body']), Paragraph('_' * 30, sty['body']), Paragraph('_' * 20, sty['body'])]]
        sig_t = Table(sig_data, colWidths=[55 * mm, 75 * mm, 40 * mm])
        sig_t.setStyle(TableStyle([('TOPPADDING', (0, 0), (-1, -1), 3), ('BOTTOMPADDING', (0, 0), (-1, -1), 3)]))
        story.append(sig_t)
    doc.build(story)
    return buf.getvalue()

def build_company_profile_pdf(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)

    def on_page(c, doc):
        pw, ph = (PAGE_W, PAGE_H)
        _draw_rect(c, 0, ph - 28 * mm, pw, 28 * mm, pri)
        _draw_rect(c, 0, ph - 30 * mm, pw, 2, sec)
        _draw_rect(c, 0, 0, pw, 12 * mm, pri)
        c.saveState()
        c.setFont('Helvetica', 7)
        c.setFillColor(HexColor('#AABBCC'))
        c.drawCentredString(pw / 2, 4 * mm, inputs.get('company_name', ''))
        c.restoreState()
    buf = io.BytesIO()
    frame = Frame(20 * mm, 18 * mm, PAGE_W - 40 * mm, PAGE_H - 52 * mm, id='main', showBoundary=0)
    pt = PageTemplate(id='main', frames=[frame], onPage=on_page)
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm, topMargin=36 * mm, bottomMargin=20 * mm)
    doc.addPageTemplates([pt])
    story = []
    story.append(Paragraph(inputs.get('company_name', ''), sty['h1']))
    if inputs.get('tagline'):
        story.append(Paragraph(inputs['tagline'], sty['accent']))
    story.append(Spacer(1, 4 * mm))
    stats = inputs.get('company_stats') or {}
    if stats:
        stat_items = [[Paragraph(f"<b>{v}</b><br/><font size='7'>{k}</font>", ParagraphStyle('st', fontSize=11, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=_hex(sec))) for k, v in stats.items()]]
        st_t = Table(stat_items, colWidths=[170 * mm / max(len(stats), 1)] * len(stats))
        st_t.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), HexColor('#F0F4F8')), ('BOX', (0, 0), (-1, -1), 0.5, HexColor('#DDDDDD')), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6)]))
        story.append(st_t)
        story.append(Spacer(1, 5 * mm))
    for section_title, key in [('About Us', 'about'), ('Our Mission', 'mission_statement'), ('Our Services', None), ('Why Choose Us', 'why_us')]:
        if key and ai_content.get(key):
            story.append(Paragraph(section_title, sty['h2']))
            story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
            story.append(Spacer(1, 2 * mm))
            for para in ai_content[key].split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), sty['body']))
                    story.append(Spacer(1, 2 * mm))
            story.append(Spacer(1, 4 * mm))
    services = ai_content.get('services', [])
    if services:
        story.append(Paragraph('Our Services', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        for svc in services:
            if ':' in svc:
                name, desc = svc.split(':', 1)
                story.append(Paragraph(f'<b>{name.strip()}</b> — {desc.strip()}', sty['body']))
            else:
                story.append(Paragraph(svc, sty['body']))
            story.append(Spacer(1, 1.5 * mm))
        story.append(Spacer(1, 4 * mm))
    team = inputs.get('team_members') or []
    if team:
        story.append(Paragraph('Our Team', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        team_rows = []
        row = []
        for tm in team:
            cell = Paragraph(f"<b>{tm.get('name', '')}</b><br/><font size='7.5'>{tm.get('title', '')}</font>", sty['body'])
            row.append(cell)
            if len(row) == 3:
                team_rows.append(row)
                row = []
        if row:
            while len(row) < 3:
                row.append(Paragraph('', sty['body']))
            team_rows.append(row)
        if team_rows:
            tt = Table(team_rows, colWidths=[56 * mm, 56 * mm, 58 * mm])
            tt.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, -1), HexColor('#F8F9FA')), ('BOX', (0, 0), (-1, -1), 0.5, HexColor('#DDDDDD')), ('TOPPADDING', (0, 0), (-1, -1), 6), ('BOTTOMPADDING', (0, 0), (-1, -1), 6), ('LEFTPADDING', (0, 0), (-1, -1), 6)]))
            story.append(tt)
        story.append(Spacer(1, 4 * mm))
    if ai_content.get('closing'):
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#CCCCCC')))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(ai_content['closing'], ParagraphStyle('closing', fontSize=10, fontName='Helvetica-BoldOblique', textColor=_hex(sec), alignment=TA_CENTER)))
    doc.build(story)
    return buf.getvalue()

def build_capability_pdf(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)

    def on_page(c, doc):
        pw, ph = (PAGE_W, PAGE_H)
        _draw_rect(c, 0, ph - 24 * mm, pw, 24 * mm, pri)
        _draw_rect(c, 0, ph - 26 * mm, pw * 0.4, 2, sec)
        c.saveState()
        c.setFillColor(HexColor('#FFFFFF'))
        c.setFont('Helvetica-Bold', 13)
        c.drawString(14 * mm, ph - 15 * mm, inputs.get('company_name', ''))
        c.setFont('Helvetica', 7.5)
        c.setFillColor(HexColor('#AABBCC'))
        meta_parts = []
        if inputs.get('duns_number'):
            meta_parts.append(f"DUNS: {inputs['duns_number']}")
        if inputs.get('cage_code'):
            meta_parts.append(f"CAGE: {inputs['cage_code']}")
        if meta_parts:
            c.drawString(14 * mm, ph - 22 * mm, '  |  '.join(meta_parts))
        c.restoreState()
        _draw_rect(c, 0, 0, pw, 10 * mm, pri)
    buf = io.BytesIO()
    frame = Frame(20 * mm, 16 * mm, PAGE_W - 40 * mm, PAGE_H - 48 * mm, id='main', showBoundary=0)
    pt = PageTemplate(id='main', frames=[frame], onPage=on_page)
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm, topMargin=32 * mm, bottomMargin=18 * mm)
    doc.addPageTemplates([pt])
    story = []
    naics = inputs.get('naics_codes') or []
    if naics:
        story.append(Paragraph('NAICS Codes', sty['h3']))
        for n in naics:
            story.append(Paragraph(f"<b>{n.get('code', '')}:</b> {n.get('description', '')}", sty['small']))
        story.append(Spacer(1, 3 * mm))
    if ai_content.get('opening'):
        story.append(Paragraph(ai_content['opening'], ParagraphStyle('opening', fontSize=12, fontName='Helvetica-Bold', textColor=_hex(pri), leading=16)))
        story.append(Spacer(1, 3 * mm))
    comp_items = ai_content.get('core_competencies', [])
    comp_paras = [Paragraph('Core Competencies', sty['h2'])]
    comp_paras.append(HRFlowable(width=85 * mm, thickness=1.5, color=_hex(sec)))
    comp_paras.append(Spacer(1, 2 * mm))
    for item in comp_items:
        if ':' in item:
            name, desc = item.split(':', 1)
            comp_paras.append(Paragraph(f'<b>✓ {name.strip()}</b>', sty['body']))
            comp_paras.append(Paragraph(desc.strip(), sty['small']))
        else:
            comp_paras.append(Paragraph(f'✓ {item}', sty['body']))
        comp_paras.append(Spacer(1, 1.5 * mm))
    diff_paras = [Paragraph('Our Differentiator', sty['h2'])]
    diff_paras.append(HRFlowable(width=75 * mm, thickness=1.5, color=_hex(sec)))
    diff_paras.append(Spacer(1, 2 * mm))
    if ai_content.get('differentiator'):
        diff_paras.append(Paragraph(ai_content['differentiator'], sty['body']))
    certs = inputs.get('certifications') or []
    if certs:
        diff_paras.append(Spacer(1, 3 * mm))
        diff_paras.append(Paragraph('Certifications', sty['h3']))
        for cert in certs:
            diff_paras.append(Paragraph(f'• {cert}', sty['small']))
    two_col = Table([[comp_paras, diff_paras]], colWidths=[92 * mm, 78 * mm])
    two_col.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 4)]))
    story.append(two_col)
    story.append(Spacer(1, 5 * mm))
    past_perf = inputs.get('past_performance') or []
    exp_highlights = ai_content.get('experience_highlights', '')
    if past_perf or exp_highlights:
        story.append(Paragraph('Past Performance', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        if exp_highlights:
            for para in exp_highlights.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), sty['body']))
                    story.append(Spacer(1, 2 * mm))
        for pp in past_perf:
            story.append(Paragraph(f"<b>{pp.get('client', '')}</b>{('(' + pp['year'] + ')' if pp.get('year') else '')}: {pp.get('description', '')}", sty['small']))
            story.append(Spacer(1, 1 * mm))
        story.append(Spacer(1, 4 * mm))
    if ai_content.get('call_to_action'):
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#CCCCCC')))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(ai_content['call_to_action'], ParagraphStyle('cta', fontSize=9, fontName='Helvetica-Bold', textColor=_hex(sec))))
    doc.build(story)
    return buf.getvalue()

def build_brand_guidelines_pdf(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color') or DEFAULT_PRIMARY
    sec = inputs.get('secondary_color') or DEFAULT_SECONDARY
    sty = _styles(pri, sec)

    def on_page(c, doc):
        pw, ph = (PAGE_W, PAGE_H)
        _draw_rect(c, 0, ph - 20 * mm, pw, 20 * mm, pri)
        _draw_rect(c, 0, ph - 22 * mm, pw, 2, sec)
        c.saveState()
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(HexColor('#FFFFFF'))
        c.drawString(14 * mm, ph - 13 * mm, f"{inputs.get('brand_name', '')} — Brand Guidelines")
        c.restoreState()
        _draw_rect(c, 0, 0, pw, 8 * mm, pri)
    buf = io.BytesIO()
    frame = Frame(20 * mm, 14 * mm, PAGE_W - 40 * mm, PAGE_H - 46 * mm, id='main', showBoundary=0)
    pt = PageTemplate(id='main', frames=[frame], onPage=on_page)
    doc = BaseDocTemplate(buf, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm, topMargin=28 * mm, bottomMargin=16 * mm)
    doc.addPageTemplates([pt])
    story = []
    story.append(Paragraph(inputs.get('brand_name', ''), sty['h1']))
    story.append(Paragraph('Brand Identity Guidelines', sty['accent']))
    story.append(Spacer(1, 6 * mm))
    sections = [('Brand Story', 'brand_story'), ('Mission', 'mission_statement'), ('Vision', 'vision_statement')]
    for title, key in sections:
        val = ai_content.get(key, '')
        if val:
            story.append(Paragraph(title, sty['h2']))
            story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
            story.append(Spacer(1, 2 * mm))
            for para in val.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), sty['body']))
                    story.append(Spacer(1, 2 * mm))
            story.append(Spacer(1, 4 * mm))
    personality = ai_content.get('brand_personality', [])
    if personality:
        story.append(Paragraph('Brand Personality', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        for trait in personality:
            story.append(Paragraph(f"<b>{trait.get('trait', '')}</b> — {trait.get('description', '')}", sty['body']))
            story.append(Spacer(1, 1.5 * mm))
        story.append(Spacer(1, 4 * mm))
    story.append(Paragraph('Colour Palette', sty['h2']))
    story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
    story.append(Spacer(1, 2 * mm))
    col_usage = ai_content.get('color_usage', {})
    color_rows = []
    for role, data in col_usage.items():
        if isinstance(data, dict):
            hex_val = data.get('hex', '')
            swatch = Paragraph(f'<font color="{hex_val}">■■■■</font> {hex_val}', ParagraphStyle('sw', fontSize=11))
            color_rows.append([swatch, Paragraph(f"<b>{role.capitalize()} — {data.get('name', '')}</b><br/><font size='8'>{data.get('usage', '')}</font>", sty['body'])])
    if color_rows:
        ct = Table(color_rows, colWidths=[35 * mm, 135 * mm])
        ct.setStyle(TableStyle([('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4), ('LINEBELOW', (0, 0), (-1, -2), 0.3, HexColor('#EEEEEE'))]))
        story.append(ct)
    story.append(Spacer(1, 4 * mm))
    typo = ai_content.get('typography', {})
    if typo:
        story.append(Paragraph('Typography', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(f"<b>Heading:</b> {typo.get('heading_font', '')}", sty['body']))
        story.append(Paragraph(f"<b>Body:</b> {typo.get('body_font', '')}", sty['body']))
        if typo.get('usage_rules'):
            story.append(Paragraph(typo['usage_rules'], sty['small']))
        story.append(Spacer(1, 4 * mm))
    tov = ai_content.get('tone_of_voice', {})
    if tov:
        story.append(Paragraph('Tone of Voice', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(tov.get('description', ''), sty['body']))
        story.append(Spacer(1, 2 * mm))
        for line in tov.get('do', []):
            story.append(Paragraph(f'✓  {line}', ParagraphStyle('do', fontSize=8, textColor=HexColor('#2E7D32'))))
        for line in tov.get('dont', []):
            story.append(Paragraph(f'✗  {line}', ParagraphStyle('dont', fontSize=8, textColor=HexColor('#C62828'))))
        story.append(Spacer(1, 4 * mm))
    logo_usage = ai_content.get('logo_usage', {})
    if logo_usage:
        story.append(Paragraph('Logo Usage', sty['h2']))
        story.append(HRFlowable(width='100%', thickness=1.5, color=_hex(sec)))
        story.append(Spacer(1, 2 * mm))
        for rule in logo_usage.get('dos', []):
            story.append(Paragraph(f'✓  {rule}', ParagraphStyle('do2', fontSize=8, textColor=HexColor('#2E7D32'))))
        for rule in logo_usage.get('donts', []):
            story.append(Paragraph(f'✗  {rule}', ParagraphStyle('dont2', fontSize=8, textColor=HexColor('#C62828'))))
    doc.build(story)
    return buf.getvalue()

def _docx_set_color(para, hex_color: str):
    for run in para.runs:
        run.font.color.rgb = RGBColor(*[int(hex_color.lstrip('#')[i:i + 2], 16) for i in (0, 2, 4)])

def _docx_heading(doc, text, level, color_hex):
    p = doc.add_heading(text, level=level)
    _docx_set_color(p, color_hex)
    return p

def _docx_add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(cell_text)
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

def _docx_shade_row(row, hex_color: str):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.lstrip('#'))
        tcPr.append(shd)

def build_letterhead_docx(inputs: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    sec = inputs.get('secondary_color', DEFAULT_SECONDARY).lstrip('#')
    doc = DocxDocument()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    hdr = section.header
    para = hdr.paragraphs[0]
    para.clear()
    run = para.add_run(inputs.get('company_name', ''))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*[int(pri[i:i + 2], 16) for i in (0, 2, 4)])
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para2 = hdr.add_paragraph()
    run2 = para2.add_run(f"  {inputs.get('phone', '')}  |  {inputs.get('email', '')}  |  {inputs.get('website', '')}")
    run2.font.size = Pt(8)
    para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_heading('Letterhead Template', level=1)
    p = doc.add_paragraph('Dear Sir / Madam,\n\nReplace this content with your correspondence.\n\nYours faithfully,')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_invoice_docx(inputs: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    sec = inputs.get('secondary_color', DEFAULT_SECONDARY).lstrip('#')
    sym = (inputs.get('currency') or 'NGN ₦').split()[-1]
    pri_rgb = RGBColor(*[int(pri[i:i + 2], 16) for i in (0, 2, 4)])
    sec_rgb = RGBColor(*[int(sec[i:i + 2], 16) for i in (0, 2, 4)])
    today = datetime.now().strftime('%d %B %Y')
    doc = DocxDocument()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    h = doc.add_heading('INVOICE', level=1)
    for run in h.runs:
        run.font.color.rgb = pri_rgb
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(inputs.get('company_name', ''))
    p.runs[0].bold = True
    doc.add_paragraph(inputs.get('company_address', ''))
    doc.add_paragraph(f"{inputs.get('phone', '')}  |  {inputs.get('email', '')}")
    doc.add_paragraph('')
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(3.2)
    tbl.columns[1].width = Inches(3.2)
    l_cell = tbl.rows[0].cells[0]
    l_cell.text = f"INVOICE TO:\n{inputs.get('client_name', '[Client Name]')}\n{inputs.get('client_address', '')}"
    r_cell = tbl.rows[0].cells[1]
    inv_prefix = inputs.get('invoice_number_prefix', 'INV-')
    r_cell.text = f"Invoice #: {inv_prefix}{inputs.get('invoice_number', '001')}\nDate: {inputs.get('invoice_date', today)}\nDue: {inputs.get('due_date', '')}\nTerms: {inputs.get('payment_terms', '')}"
    doc.add_paragraph('')
    items = inputs.get('items', [])
    items_tbl = doc.add_table(rows=1, cols=5)
    items_tbl.style = 'Table Grid'
    hrow = items_tbl.rows[0]
    for i, h in enumerate(['#', 'Description', 'Qty', 'Unit Price', 'Total']):
        hrow.cells[i].text = h
        for run in hrow.cells[i].paragraphs[0].runs:
            run.bold = True
    _docx_shade_row(hrow, pri)
    subtotal = 0.0
    for i, item in enumerate(items):
        qty = item.get('qty', item.get('quantity', 1))
        price = float(item.get('unit_price', item.get('price', 0)))
        total = float(item.get('total', qty * price))
        subtotal += total
        row = items_tbl.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = item.get('description', '')
        row.cells[2].text = str(qty)
        row.cells[3].text = f'{sym}{price:,.2f}'
        row.cells[4].text = f'{sym}{total:,.2f}'
    doc.add_paragraph('')
    tax_rate = float(inputs.get('tax_rate', 0))
    discount = inputs.get('discount', 0) or 0
    taxable = subtotal - discount
    tax_amt = taxable * tax_rate / 100
    grand = taxable + tax_amt
    totals_tbl = doc.add_table(rows=0, cols=2)
    totals_rows = [('Subtotal', f'{sym}{subtotal:,.2f}')]
    if discount:
        totals_rows.append(('Discount', f'-{sym}{discount:,.2f}'))
    if tax_rate:
        totals_rows.append((f'Tax ({tax_rate:.1f}%)', f'{sym}{tax_amt:,.2f}'))
    totals_rows.append(('TOTAL', f'{sym}{grand:,.2f}'))
    for lbl, val in totals_rows:
        row = totals_tbl.add_row()
        row.cells[0].text = lbl
        row.cells[1].text = val
        if lbl == 'TOTAL':
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = pri_rgb
    pm_list = inputs.get('payment_methods') or []
    if pm_list:
        doc.add_paragraph('')
        p = doc.add_paragraph('PAYMENT METHODS')
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = pri_rgb
        for pm in pm_list:
            doc.add_paragraph(f"{pm.get('label', '')}: {pm.get('bank_name', '')} | A/C: {pm.get('account_number', '')} | {pm.get('account_name', '')}")
    if inputs.get('footer_note'):
        doc.add_paragraph('')
        doc.add_paragraph(f"Notes: {inputs['footer_note']}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_quotation_docx(inputs: dict, asset_id: str='default') -> bytes:
    pri = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    sym = (inputs.get('currency') or 'NGN ₦').split()[-1]
    pri_rgb = RGBColor(*[int(pri[i:i + 2], 16) for i in (0, 2, 4)])
    today = datetime.now().strftime('%d %B %Y')
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.8)
    h = doc.add_heading('QUOTATION', level=1)
    for run in h.runs:
        run.font.color.rgb = pri_rgb
    doc.add_paragraph(f"Company: {inputs.get('company_name', '')}")
    doc.add_paragraph(f"Address: {inputs.get('company_address', '')}")
    if inputs.get('registration_number'):
        doc.add_paragraph(f"Reg No: {inputs['registration_number']}")
    doc.add_paragraph(f"Quote #: {inputs.get('quote_number_prefix', 'QT-')}{inputs.get('quote_number', '001')}  |  Date: {today}  |  Valid: {inputs.get('quote_valid_for', '30 days')}")
    if inputs.get('expiration_date'):
        doc.add_paragraph(f"Expires: {inputs['expiration_date']}")
    if inputs.get('prepared_by'):
        doc.add_paragraph(f"Prepared By: {inputs['prepared_by']}")
    doc.add_paragraph('')
    h2 = doc.add_paragraph('Bill To:')
    h2.runs[0].bold = True
    doc.add_paragraph(inputs.get('client_name', '[Client Name]'))
    doc.add_paragraph(inputs.get('client_address', ''))
    doc.add_paragraph('')
    items = inputs.get('items', [])
    items_tbl = doc.add_table(rows=1, cols=5)
    items_tbl.style = 'Table Grid'
    hrow = items_tbl.rows[0]
    for i, hdr in enumerate(['#', 'Description', 'Qty', 'Unit Price', 'Total']):
        hrow.cells[i].text = hdr
        for run in hrow.cells[i].paragraphs[0].runs:
            run.bold = True
    _docx_shade_row(hrow, pri)
    subtotal = 0.0
    for i, item in enumerate(items):
        qty = item.get('qty', item.get('quantity', 1))
        price = float(item.get('unit_price', item.get('price', 0)))
        total = float(item.get('total', qty * price))
        subtotal += total
        row = items_tbl.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = item.get('description', '')
        row.cells[2].text = str(qty)
        row.cells[3].text = f'{sym}{price:,.2f}'
        row.cells[4].text = f'{sym}{total:,.2f}'
    doc.add_paragraph('')
    tax_rate = float(inputs.get('tax_rate', inputs.get('vat_rate', 0)))
    grand = subtotal * (1 + tax_rate / 100)
    tot = doc.add_paragraph(f'Grand Total: {sym}{grand:,.2f}')
    tot.runs[0].bold = True
    tot.runs[0].font.color.rgb = pri_rgb
    if inputs.get('terms_and_conditions'):
        doc.add_paragraph('')
        p = doc.add_paragraph('Terms and Conditions')
        p.runs[0].bold = True
        p.runs[0].font.color.rgb = pri_rgb
        doc.add_paragraph(inputs['terms_and_conditions'])
    if inputs.get('signature_section', True):
        doc.add_paragraph('')
        doc.add_paragraph('Name: _________________________    Signature: _________________________    Date: ____________')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_company_profile_docx(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri_hex = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    pri_rgb = RGBColor(*[int(pri_hex[i:i + 2], 16) for i in (0, 2, 4)])
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    h = doc.add_heading(inputs.get('company_name', ''), level=1)
    for run in h.runs:
        run.font.color.rgb = pri_rgb
    if inputs.get('tagline'):
        doc.add_paragraph(inputs['tagline'])
    for title, key in [('About Us', 'about'), ('Mission', 'mission_statement'), ('Why Choose Us', 'why_us'), ('Closing', 'closing')]:
        val = ai_content.get(key)
        if val:
            doc.add_heading(title, level=2)
            doc.add_paragraph(val)
    services = ai_content.get('services', [])
    if services:
        doc.add_heading('Our Services', level=2)
        for s in services:
            doc.add_paragraph(s, style='List Bullet')
    team = inputs.get('team_members') or []
    if team:
        doc.add_heading('Our Team', level=2)
        for tm in team:
            p = doc.add_paragraph(f"{tm.get('name', '')} — {tm.get('title', '')}")
            p.runs[0].bold = True
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_capability_docx(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri_hex = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    pri_rgb = RGBColor(*[int(pri_hex[i:i + 2], 16) for i in (0, 2, 4)])
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    h = doc.add_heading(inputs.get('company_name', ''), level=1)
    for run in h.runs:
        run.font.color.rgb = pri_rgb
    doc.add_heading('Capability Statement', level=2)
    if ai_content.get('opening'):
        doc.add_paragraph(ai_content['opening']).runs[0].bold = True
    naics = inputs.get('naics_codes') or []
    if naics:
        doc.add_heading('NAICS Codes', level=3)
        for n in naics:
            doc.add_paragraph(f"{n.get('code', '')}: {n.get('description', '')}", style='List Bullet')
    comps = ai_content.get('core_competencies', [])
    if comps:
        doc.add_heading('Core Competencies', level=2)
        for c in comps:
            doc.add_paragraph(c, style='List Bullet')
    if ai_content.get('differentiator'):
        doc.add_heading('Our Differentiator', level=2)
        doc.add_paragraph(ai_content['differentiator'])
    if ai_content.get('experience_highlights'):
        doc.add_heading('Past Performance', level=2)
        doc.add_paragraph(ai_content['experience_highlights'])
    certs = inputs.get('certifications') or []
    if certs:
        doc.add_heading('Certifications', level=2)
        for c in certs:
            doc.add_paragraph(c, style='List Bullet')
    if ai_content.get('call_to_action'):
        doc.add_paragraph(ai_content['call_to_action'])
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_brand_guidelines_docx(inputs: dict, ai_content: dict, asset_id: str='default') -> bytes:
    pri_hex = inputs.get('primary_color', DEFAULT_PRIMARY).lstrip('#')
    pri_rgb = RGBColor(*[int(pri_hex[i:i + 2], 16) for i in (0, 2, 4)])
    doc = DocxDocument()
    h = doc.add_heading(f"{inputs.get('brand_name', '')} — Brand Guidelines", level=1)
    for run in h.runs:
        run.font.color.rgb = pri_rgb
    for title, key in [('Brand Story', 'brand_story'), ('Mission', 'mission_statement'), ('Vision', 'vision_statement')]:
        val = ai_content.get(key)
        if val:
            doc.add_heading(title, level=2)
            doc.add_paragraph(val)
    personality = ai_content.get('brand_personality', [])
    if personality:
        doc.add_heading('Brand Personality', level=2)
        for trait in personality:
            p = doc.add_paragraph(f"{trait.get('trait', '')}: {trait.get('description', '')}")
            p.runs[0].bold = True
    col_usage = ai_content.get('color_usage', {})
    if col_usage:
        doc.add_heading('Colour Palette', level=2)
        for role, data in col_usage.items():
            if isinstance(data, dict):
                doc.add_paragraph(f"{role}: {data.get('hex', '')} — {data.get('usage', '')}")
    typo = ai_content.get('typography', {})
    if typo:
        doc.add_heading('Typography', level=2)
        doc.add_paragraph(f"Heading: {typo.get('heading_font', '')}")
        doc.add_paragraph(f"Body: {typo.get('body_font', '')}")
        doc.add_paragraph(typo.get('usage_rules', ''))
    tov = ai_content.get('tone_of_voice', {})
    if tov:
        doc.add_heading('Tone of Voice', level=2)
        doc.add_paragraph(tov.get('description', ''))
        for d in tov.get('do', []):
            doc.add_paragraph(f'✓ {d}', style='List Bullet')
        for d in tov.get('dont', []):
            doc.add_paragraph(f'✗ {d}', style='List Bullet')
    logo_usage = ai_content.get('logo_usage', {})
    if logo_usage:
        doc.add_heading('Logo Usage', level=2)
        for rule in logo_usage.get('dos', []):
            doc.add_paragraph(f'✓ {rule}', style='List Bullet')
        for rule in logo_usage.get('donts', []):
            doc.add_paragraph(f'✗ {rule}', style='List Bullet')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

class DocumentService:

    def build_letterhead_pdf(self, inputs, asset_id='default'):
        return build_letterhead_pdf(inputs, asset_id)

    def build_letterhead_docx(self, inputs, asset_id='default'):
        return build_letterhead_docx(inputs, asset_id)

    def build_invoice_pdf(self, inputs, asset_id='default'):
        return build_invoice_pdf(inputs, asset_id)

    def build_invoice_docx(self, inputs, asset_id='default'):
        return build_invoice_docx(inputs, asset_id)

    def build_quotation_pdf(self, inputs, asset_id='default'):
        return build_quotation_pdf(inputs, asset_id)

    def build_quotation_docx(self, inputs, asset_id='default'):
        return build_quotation_docx(inputs, asset_id)

    def build_company_profile_pdf(self, inputs, ai, asset_id='default'):
        return build_company_profile_pdf(inputs, ai, asset_id)

    def build_company_profile_docx(self, inputs, ai, asset_id='default'):
        return build_company_profile_docx(inputs, ai, asset_id)

    def build_capability_pdf(self, inputs, ai, asset_id='default'):
        return build_capability_pdf(inputs, ai, asset_id)

    def build_capability_docx(self, inputs, ai, asset_id='default'):
        return build_capability_docx(inputs, ai, asset_id)

    def build_brand_guidelines_pdf(self, inputs, ai, asset_id='default'):
        return build_brand_guidelines_pdf(inputs, ai, asset_id)

    def build_brand_guidelines_docx(self, inputs, ai, asset_id='default'):
        return build_brand_guidelines_docx(inputs, ai, asset_id)
document_service = DocumentService()